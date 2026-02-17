import streamlit as st
import json
import time
import re
import logging
from typing import Dict, Any, List, Optional, Generator, Tuple
from datetime import datetime
from document_generator import DocumentGenerator
from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from database import DatabaseManager
from modal_close_helper import confirm_generation_started, is_waiting_for_generation

# Configure logging
logger = logging.getLogger(__name__)

# Autosave interval
AUTOSAVE_INTERVAL_SECONDS = 30

def get_document_id() -> Optional[str]:
    """Retrieves the current document ID from URL or session state."""
    # 1. Check URL parameters (highest priority)
    try:
        doc_id = st.query_params.get("document_id")
        if doc_id:
            return str(doc_id)
    except Exception:
        pass
        
    # 2. Check current session
    if "current_document_id" in st.session_state and st.session_state.current_document_id:
        return str(st.session_state.current_document_id)
    
    # 3. Check preserved ID (fallback)
    if "preserved_document_id" in st.session_state and st.session_state.preserved_document_id:
        return str(st.session_state.preserved_document_id)
        
    return None

def _render_loading_animation():
    """Render a clean loading spinner for the document editor."""
    st.markdown("""
    <style>
    @keyframes pulse {
        0%, 100% { opacity: 0.4; }
        50% { opacity: 1; }
    }
    
    @keyframes slideUp {
        from {
            opacity: 0;
            transform: translateY(10px);
        }
        to {
            opacity: 1;
            transform: translateY(0);
        }
    }
    
    .doc-loading-container {
        display: flex;
        flex-direction: column;
        align-items: center;
        justify-content: center;
        min-height: 70vh;
        padding: 40px;
        animation: slideUp 0.4s ease-out;
    }
    
    .doc-loading-spinner {
        width: 60px;
        height: 60px;
        border: 4px solid #252930;
        border-top: 4px solid #4B9EFF;
        border-radius: 50%;
        animation: spin 1s linear infinite;
        margin-bottom: 24px;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    .doc-loading-text {
        font-size: 18px;
        font-weight: 600;
        color: #FFFFFF;
        margin-bottom: 8px;
    }
    
    .doc-loading-subtext {
        font-size: 14px;
        color: #9BA1B0;
        animation: pulse 2s ease-in-out infinite;
    }
    </style>
    
    <div class="doc-loading-container">
        <div class="doc-loading-spinner"></div>
        <div class="doc-loading-text">Loading Document</div>
        <div class="doc-loading-subtext">Preparing your document for editing...</div>
    </div>
    """, unsafe_allow_html=True)

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def _word_count(html_text: str) -> int:
    """Extract plain text and count words."""
    soup = BeautifulSoup(html_text or "", 'html.parser')
    text = soup.get_text()
    words = re.findall(r"\b[\w'''-]+\b", text)
    return len(words)


def _extract_plain_text(html_text: str) -> str:
    """Extract plain text from HTML."""
    soup = BeautifulSoup(html_text or "", 'html.parser')
    return soup.get_text()


def _extract_headings(html_text: str) -> List[Dict[str, Any]]:
    """Extract headings for outline generation."""
    soup = BeautifulSoup(html_text or "", 'html.parser')
    headings = []
    
    for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
        heading_id = tag.get('id')
        
        if heading_id:
            level = int(tag.name[1])
            headings.append({
                'level': level,
                'text': tag.get_text().strip(),
                'id': heading_id 
            })
            
    return headings


def _sanitize_html(html: str) -> str:
    """Basic HTML sanitization to prevent XSS."""
    soup = BeautifulSoup(html, 'html.parser')
    for tag in soup.find_all('script'):
        tag.decompose()
    for tag in soup.find_all():
        if 'onclick' in tag.attrs:
            del tag.attrs['onclick']
        if 'onerror' in tag.attrs:
            del tag.attrs['onerror']
    return str(soup)


def _html_to_docx(html_content: str, docx_path: str, matter_name: str):
    """Convert HTML to DOCX preserving editor formatting + applying legal standards."""
    doc = DocxDocument()
    
    # --- 1. Setup Base Styles (Legal Standard) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.15  # Standard legal spacing
    paragraph_format.space_after = Pt(12)

    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process elements
    for element in soup.children:
        _process_legal_element(element, doc)
    
    doc.save(docx_path)

def _parse_css_style(style_str: str) -> dict:
    """Parses an inline style string into a dictionary."""
    if not style_str:
        return {}
    styles = {}
    for item in style_str.split(';'):
        if ':' in item:
            key, value = item.split(':', 1)
            styles[key.strip().lower()] = value.strip().lower()
    return styles

def _px_to_inches(value_str: str) -> float:
    """Converts px, em, or pt strings to Inches float."""
    try:
        if 'px' in value_str:
            return float(value_str.replace('px', '')) / 96.0  # Standard web DPI
        if 'em' in value_str:
            return float(value_str.replace('em', '')) * 0.17  # Approx 1em = 12pt = 0.17in
        if 'pt' in value_str:
            return float(value_str.replace('pt', '')) / 72.0
        if value_str.replace('.', '', 1).isdigit():
            return float(value_str) # Assume inches if no unit
    except:
        return 0.0
    return 0.0

def _process_legal_element(element, doc, parent_styles=None):
    """
    Smart processor that honors User CSS first, then falls back to Legal Heuristics.
    """
    # 1. Handle Text Nodes
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            p = doc.add_paragraph()
            p.add_run(text)
        return

    if not hasattr(element, 'name') or not element.name:
        return

    # Extract CSS styles from this element
    css_styles = _parse_css_style(element.get('style', ''))
    
    # Merge with parent styles if needed (simple inheritance)
    if parent_styles:
        # We prioritize current element styles over parent
        css_styles = {**parent_styles, **css_styles}

    # 2. Heuristic Detection (Backup Logic)
    text_content = element.get_text().strip()
    is_legal_header = False
    header_patterns = [
        r"^REPUBLIC OF KENYA$", r"^IN THE MATTER OF", r"^AFFIDAVIT$", 
        r"^TENANCY AGREEMENT$", r"^THIS TENANCY AGREEMENT"
    ]
    if any(re.match(pattern, text_content, re.IGNORECASE) for pattern in header_patterns):
        is_legal_header = True

    # 3. Handle Paragraphs & Headings
    if element.name in ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        p = doc.add_paragraph()
        
        # --- A. APPLY ALIGNMENT (User > Heuristic) ---
        align_map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            'left': WD_ALIGN_PARAGRAPH.LEFT
        }
        
        if 'text-align' in css_styles:
            # User explicitly set alignment
            p.alignment = align_map.get(css_styles['text-align'], WD_ALIGN_PARAGRAPH.LEFT)
        elif is_legal_header or element.name == 'h1':
            # Fallback: Auto-center legal headers
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # --- B. APPLY INDENTATION (User > Heuristic) ---
        if 'margin-left' in css_styles or 'padding-left' in css_styles:
            val = css_styles.get('margin-left') or css_styles.get('padding-left')
            p.paragraph_format.left_indent = Inches(_px_to_inches(val))
        elif text_content.startswith("THAT ") or text_content.startswith("1. THAT"):
            # Fallback: Auto-indent "THAT" clauses
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

        # --- C. APPLY LINE HEIGHT ---
        if 'line-height' in css_styles:
            try:
                # If unitless (e.g., "1.5"), use as multiple. If "px", convert? 
                # Word handles line_spacing as float multiplier usually
                val = css_styles['line-height']
                if 'px' not in val:
                    p.paragraph_format.line_spacing = float(val)
            except:
                pass

        # Populate text
        _add_formatted_text_to_paragraph(element, p, inherited_styles=css_styles)

    # 4. Handle Lists 
    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            # Apply list styles if present on the UL/LI
            li_styles = _parse_css_style(li.get('style', ''))
            full_styles = {**css_styles, **li_styles}
            _add_formatted_text_to_paragraph(li, p, inherited_styles=full_styles)
    
    elif element.name == 'ol':
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Number')
            li_styles = _parse_css_style(li.get('style', ''))
            full_styles = {**css_styles, **li_styles}
            _add_formatted_text_to_paragraph(li, p, inherited_styles=full_styles)

    # 5. Handle Tables
    elif element.name == 'table':
        rows = element.find_all('tr')
        if rows:
            cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            
            # Use 'Table Grid' so user sees borders, unless explicitly hidden
            table.style = 'Table Grid'
            if 'border' in css_styles and 'none' in css_styles['border']:
                table.style = 'Normal Table' # Invisible borders

            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    if j < len(table.rows[i].cells):
                        cell_p = table.rows[i].cells[j].paragraphs[0]
                        _add_formatted_text_to_paragraph(cell, cell_p)
                        
                        # Auto-bold signatures if valid
                        cell_text = cell.get_text().strip()
                        if "Signed by" in cell_text or "In the presence of" in cell_text:
                            for run in cell_p.runs:
                                run.bold = True

    # 6. Blockquotes
    elif element.name == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        _add_formatted_text_to_paragraph(element, p)
        
    # 7. Fallback for containers
    elif hasattr(element, 'children'):
        for child in element.children:
            _process_legal_element(child, doc, parent_styles=css_styles)

def _add_formatted_text_to_paragraph(element: Tag, paragraph, inherited_styles=None):
    """
    Parses HTML tags AND CSS styles to apply formatting to Runs.
    """
    def process_node(node, current_styles=None):
        if current_styles is None:
            current_styles = {}
        
        # Merge inherited styles (from p or div) with current node styles
        node_styles = {}
        if isinstance(node, Tag):
            node_styles = _parse_css_style(node.get('style', ''))
        
        # Effective styles for this node
        active_styles = {**current_styles, **node_styles}

        if isinstance(node, str):
            text = node
            if not text: return
            
            # Split for smart "THAT" handling
            parts = re.split(r'(THAT\s)', text) 
            
            for part in parts:
                if not part: continue
                
                run = paragraph.add_run(part)
                
                # --- APPLY STYLES ---
                # 1. CSS-based Bold
                if ('font-weight' in active_styles and 
                    ('bold' in active_styles['font-weight'] or '700' in active_styles['font-weight'])):
                    run.bold = True
                
                # 2. CSS-based Italic
                if 'font-style' in active_styles and 'italic' in active_styles['font-style']:
                    run.italic = True
                
                # 3. CSS-based Underline
                if 'text-decoration' in active_styles and 'underline' in active_styles['text-decoration']:
                    run.underline = True

                # 4. CSS-based Font Size
                if 'font-size' in active_styles:
                     # Attempt to parse "14pt" or "18px"
                     val = active_styles['font-size']
                     if 'pt' in val:
                         run.font.size = Pt(float(val.replace('pt','')))
                     elif 'px' in val:
                         # 1px approx 0.75pt
                         run.font.size = Pt(float(val.replace('px','')) * 0.75)

                # --- APPLY HEURISTICS (Only if plain text) ---
                if part.strip() == "THAT":
                    run.bold = True
                    run.underline = True
                
                if re.match(r'KES\s?[\d,]+', part):
                    run.bold = True

        elif isinstance(node, Tag):
            if node.name == 'br':
                paragraph.add_run('\n')
                return
            
            # Map HTML tags to styles
            tag_styles = active_styles.copy()
            if node.name in ['strong', 'b']: tag_styles['font-weight'] = 'bold'
            if node.name in ['em', 'i']: tag_styles['font-style'] = 'italic'
            if node.name == 'u': tag_styles['text-decoration'] = 'underline'
            
            for child in node.children:
                process_node(child, tag_styles)

    if inherited_styles:
        # Pass paragraph-level styles down to text runs if relevant (like color/font)
        # though usually bold/italic are inline.
        pass
        
    for child in element.children:
        process_node(child, inherited_styles)


def _format_datetime(dt) -> str:
    """Format datetime for display."""
    if not dt:
        return "Unknown"
    try:
        if isinstance(dt, str):
            dt = datetime.fromisoformat(dt.replace('Z', '+00:00'))
        return dt.strftime("%b %d, %Y at %I:%M %p")
    except Exception:
        return "Unknown"


def _time_ago(dt) -> str:
    """Get relative time string."""
    if not dt:
        return "Unknown"
    try:
        if isinstance(dt, str):
            dt = datetime.fromisoformat(dt.replace('Z', '+00:00'))
        now = datetime.now(dt.tzinfo) if dt.tzinfo else datetime.now()
        diff = now - dt
        
        if diff.days > 0:
            return f"{diff.days}d ago"
        elif diff.seconds >= 3600:
            hours = diff.seconds // 3600
            return f"{hours}h ago"
        elif diff.seconds >= 60:
            minutes = diff.seconds // 60
            return f"{minutes}m ago"
        else:
            return "Just now"
    except Exception:
        return "Unknown"


def _get_clauses_for_editor(db: DatabaseManager, light_mode: bool = True) -> list:
    """Fetch all clauses from the database and format them for the editor component.
    
    Args:
        db: Database manager instance
        light_mode: If True, fetch only metadata (id, title, category, tags) for faster loading.
                   If False, fetch full content (for insertion).
    """
    try:
        all_clauses = db.get_clauses(include_system=True, light_mode=light_mode)
        
        formatted_clauses = []
        for clause in all_clauses:
            if light_mode:
                # Light mode: metadata only for display in clause browser
                formatted_clauses.append({
                    'id': clause['id'],
                    'title': clause['title'],
                    'category': clause['category'],
                    'tags': clause.get('tags', []),
                    'usage_count': clause.get('usage_count', 0),
                    'is_pinned': clause.get('is_pinned', False),
                    'is_system': clause.get('is_system', False),
                    'preview': clause.get('preview', '')
                })
            else:
                # Full mode: include content for insertion
                formatted_clauses.append({
                    'id': clause['id'],
                    'title': clause['title'],
                    'category': clause['category'],
                    'content': clause['content'],
                    'tags': clause.get('tags', []),
                    'usage_count': clause.get('usage_count', 0),
                    'is_pinned': clause.get('is_pinned', False),
                    'is_system': clause.get('is_system', False)
                })
        
        return formatted_clauses
    except Exception as e:
        logger.error(f"Error fetching clauses: {e}")
        return []


def _get_comments_for_version(db: DatabaseManager, version_id: str) -> list:
    """Fetch all comments for a version from the database."""
    try:
        comments = db.get_comments(version_id, include_resolved=True)
        
        formatted_comments = []
        
        for comment in comments:
            # Use stored author info if available, otherwise fallback
            user_email = comment.get('author_email', 'Unknown User')
            user_name = comment.get('author_name', user_email.split('@')[0].title() if user_email != 'Unknown User' else 'Unknown User')
            
            formatted_comments.append({
                'id': comment['id'],
                'text': comment['comment_text'],
                'selectedText': comment.get('selected_text', ''),
                'timestamp': comment['created_at'],
                'author': user_name,
                'authorEmail': user_email,
                'resolved': comment.get('is_resolved', False)
            })
        
        logger.info(f"✅ Loaded {len(formatted_comments)} comments for version {version_id}")
        return formatted_comments
    except Exception as e:
        logger.error(f"❌ Error fetching comments: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return []

def _save_comment_to_db(db: DatabaseManager, version_id: str, comment_data: dict) -> bool:
    """Save a comment to the database."""
    try:
        logger.info(f"💾 Saving comment to database for version {version_id}")
        logger.debug(f"Comment data: {comment_data}")
        
        # Validate required fields
        if not comment_data.get('text'):
            logger.error("Comment text is required")
            return False
        
        if not version_id:
            logger.error("Version ID is required")
            return False
        
        # Create the comment
        result = db.create_comment(
            version_id=version_id,
            comment_text=comment_data['text'],
            selected_text=comment_data.get('selectedText', ''),
            position_start=comment_data.get('positionStart'),  # May be None
            position_end=comment_data.get('positionEnd')  # May be None
        )
        
        if result and result.get('id'):
            logger.info(f"✅ Comment saved successfully with ID: {result['id']}")
            return True
        else:
            logger.error("Comment creation returned no ID")
            return False
            
    except Exception as e:
        logger.error(f"❌ Error saving comment to database: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return False


# ============================================================================
# SESSION STATE INITIALIZATION
# ============================================================================

def _init_editor_state():
    """Initialize all editor session state variables."""
    if "editor_content" not in st.session_state:
        st.session_state.editor_content = ""
    
    if "editor_original_content" not in st.session_state:
        st.session_state.editor_original_content = ""
    
    if "generation_complete" not in st.session_state:
        st.session_state.generation_complete = False
    
    if "last_autosave" not in st.session_state:
        st.session_state.last_autosave = 0
    
    if "show_versions_panel" not in st.session_state:
        st.session_state.show_versions_panel = False
    
    if "unsaved_changes" not in st.session_state:
        st.session_state.unsaved_changes = False
    
    if "autosave_in_progress" not in st.session_state:
        st.session_state.autosave_in_progress = False
    
    if "last_component_content" not in st.session_state:
        st.session_state.last_component_content = ""
    
    if "pending_autosave_content" not in st.session_state:
        st.session_state.pending_autosave_content = None
    
    if "editor_comments" not in st.session_state:
        st.session_state.editor_comments = []


# ============================================================================
# PROGRESSIVE AUTOSAVE FUNCTION
# ============================================================================

def _handle_progressive_save(db: DatabaseManager, document_id: str, content: str):
    """
    Saves the document content to the current version row for autosave.
    Does NOT create a new version and does NOT trigger a rerun.
    CRITICAL FIX: This now only saves to DB without triggering any state updates.
    """
    try:
        if st.session_state.autosave_in_progress:
            print("Autosave already in progress, skipping...")
            return
        
        st.session_state.autosave_in_progress = True
        
        content_plain = _extract_plain_text(content)
        word_count = _word_count(content)
        
        if word_count < 1:
            print("Autosave skipped: no content")
            st.session_state.autosave_in_progress = False
            return

        current_version_id = st.session_state.current_version_id
        
        print(f"💾 Autosaving {word_count} words to version {current_version_id}")
        
        updated_version = db.update_version_content(
            version_id=current_version_id,
            content=content,
            content_plain=content_plain
        )
        
        if not updated_version:
            print(f"Failed to autosave (update) version {current_version_id}")
            st.session_state.autosave_in_progress = False
            return

        st.session_state.last_autosave = time.time()
        st.session_state.autosave_in_progress = False
        
        # CRITICAL FIX: Update session state silently - DO NOT modify editor_content
        # The component owns the content and will send updates when ready
        print(f"✅ Autosaved to database: {current_version_id} ({word_count} words)")

    except Exception as e:
        print(f"❌ Error during autosave: {e}")
        import traceback
        print(traceback.format_exc())
        st.session_state.autosave_in_progress = False


# ============================================================================
# COMMENT HANDLING
# ============================================================================

def _handle_comment_action(db: DatabaseManager, version_id: str, action: dict):
    """Handle comment-related actions from the component."""
    try:
        action_type = action.get('type')
        print(f"🔧 Processing comment action: {action_type}")
        
        if action_type == 'create':
            # Create new comment in database
            comment_data = action.get('comment')
            if not comment_data:
                print("❌ No comment data provided")
                return False
                
            print(f"📝 Creating comment: {comment_data.get('text', '')[:50]}...")
            success = _save_comment_to_db(db, version_id, comment_data)
            if success:
                print(f"✅ Comment saved to database successfully")
                # Reload comments from database
                st.session_state.editor_comments = _get_comments_for_version(db, version_id)
                return True
            else:
                print("❌ Failed to save comment to database")
                return False
                
        elif action_type == 'resolve':
            # Resolve comment in database
            comment_id = action.get('commentId')
            if not comment_id:
                print("❌ No comment ID provided for resolve action")
                return False
                
            print(f"✅ Resolving comment {comment_id}")
            success = db.resolve_comment(comment_id)
            if success:
                print(f"✅ Comment {comment_id} resolved successfully")
                st.session_state.editor_comments = _get_comments_for_version(db, version_id)
                return True
            else:
                print(f"❌ Failed to resolve comment {comment_id}")
                return False
                
        elif action_type == 'delete':
            # Delete comment from database
            comment_id = action.get('commentId')
            if not comment_id:
                print("❌ No comment ID provided for delete action")
                return False
                
            print(f"🗑️ Deleting comment {comment_id}")
            success = db.delete_comment(comment_id, hard_delete=False)
            if success:
                print(f"✅ Comment {comment_id} deleted successfully")
                st.session_state.editor_comments = _get_comments_for_version(db, version_id)
                return True
            else:
                print(f"❌ Failed to delete comment {comment_id}")
                return False
        
        print(f"❌ Unknown action type: {action_type}")
        return False
        
    except Exception as e:
        print(f"❌ Error handling comment action: {e}")
        import traceback
        print(traceback.format_exc())
        return False


# ============================================================================
# VERSIONS PANEL
# ============================================================================

def _render_versions_panel(db: DatabaseManager, document_id: str, current_version_id: str):
    """Render the versions management panel using the new component."""
    
    from versions_panel_component import st_versions_panel

    versions = db.get_versions(document_id, include_content=False)
  
    if not versions:
        st.markdown("""
        <div style="padding: 0 20px 20px 20px;">
            <div style="background: rgba(255, 255, 255, 0.03); border: 1px solid #252930; border-radius: 8px; padding: 24px; text-align: center;">
                <div style="font-size: 14px; color: #9BA1B0;">No versions saved yet</div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        return

    # Format versions for the component - it expects 'timestamp', not 'created_at'
    formatted_versions = []
    for version in versions:
        formatted_version = {
            'id': version['id'],
            'label': version.get('label', f"Version {version.get('version_number', '?')}"),
            'timestamp': version.get('created_at', datetime.now().isoformat()),  # Map created_at to timestamp
            'is_major_version': version.get('is_major_version', False),
            'change_summary': version.get('change_summary')
        }
        formatted_versions.append(formatted_version)

    selected_version_id = st_versions_panel(
        versions=formatted_versions,
        current_version_id=current_version_id,
        key="versions_panel_component",
        height=780
    )
    
    if selected_version_id and selected_version_id != current_version_id:
        full_version = db.get_version(selected_version_id)
        if full_version:
            st.session_state.editor_content = full_version['content']
            st.session_state.current_version_id = full_version['id']
            st.session_state.editor_original_content = full_version['content']
            st.session_state.unsaved_changes = False
            st.session_state.last_component_content = full_version['content']
            
            if not st.session_state.editor_comments or st.session_state.get("refresh_comments", False):
                st.session_state.editor_comments = _get_comments_for_version(db, full_version['id'])
                st.session_state.refresh_comments = False

            st.success(f"✅ Loaded {full_version.get('label', 'version')}")
            time.sleep(0.5)
            st.rerun()
        else:
            st.error("Failed to load version")


# ============================================================================
# EDITOR FRAGMENT - THIS PREVENTS RELOADING DURING AUTOSAVE
# ============================================================================

@st.fragment
def _render_editor_fragment(db: DatabaseManager, document_id: str, clauses_data: list, version_id: str):
    """
    Render the editor as a fragment so it doesn't reload during autosaves.
    CRITICAL FIX: Component is the source of truth for content during editing.
    Python only receives updates and persists to DB.
    """
    from document_editor_component import st_doc_editor
    
    # CRITICAL: Only use editor_content for INITIAL load
    # After that, component owns the content
    if "editor_initialized" not in st.session_state:
        st.session_state.editor_initialized = True
        current_content = st.session_state.editor_content
    else:
        # Use last known content from component, not session state
        current_content = st.session_state.get("last_component_content", st.session_state.editor_content)
    
    # Pass user information to the component
    user_email = st.session_state.get("user_email", "Unknown User")
    user_name = user_email.split('@')[0].title()
    
    # OPTIMIZATION: Defer comment loading to background after editor displays
    # This prevents blocking the initial editor render
    if not st.session_state.editor_comments and version_id:
        # Load comments asynchronously - editor will update when available
        try:
            st.session_state.editor_comments = _get_comments_for_version(db, version_id)
            logger.info(f"✅ Loaded {len(st.session_state.editor_comments)} comments in background")
        except Exception as e:
            logger.error(f"Failed to load comments: {e}")
            st.session_state.editor_comments = []
    
    # Inject user info and comments into the page for the component to access
    comments_json = json.dumps(st.session_state.editor_comments)
    comments_json_escaped = comments_json.replace('\\', '\\\\').replace('"', '\\"').replace('\n', '\\n')
    st.markdown(f"""
    <script>
        (function() {{
            console.log("🚀 Initializing comment system...");
            
            window.streamlitUserEmail = "{user_email}";
            window.streamlitUserName = "{user_name}";
            window.streamlitVersionId = "{version_id}";
            
            try {{
                const commentsData = "{comments_json_escaped}";
                window.streamlitComments = JSON.parse(commentsData.replace(/\\\\"/g, '"').replace(/\\\\n/g, '\\n'));
                console.log("✅ Loaded " + window.streamlitComments.length + " comments");
            }} catch(e) {{
                console.error("❌ Failed to parse comments JSON:", e);
                window.streamlitComments = [];
            }}
            
            window.reloadComments = function() {{
                console.log("🔄 Reloading comments from Python...");
            }};
            
            window.handleCommentAction = function(action) {{
                console.log("🔨 Comment action received from component:", action);
                
                const messageData = {{
                    type: 'streamlit:setComponentValue',
                    value: {{
                        type: 'comment_action',
                        action: action,
                        timestamp: Date.now()
                    }}
                }};
                
                if (window.parent) {{
                    console.log("📤 Sending to Streamlit parent:", messageData);
                    window.parent.postMessage(messageData, '*');
                }}
            }};
            
            console.log("✅ Comment system initialized");
        }})();
    </script>
    """, unsafe_allow_html=True)
    
    # Render the editor with current content
    result = st_doc_editor(
        content=current_content,
        clauses=clauses_data,
        comments=st.session_state.editor_comments,
        versionId=version_id,
        debounce=1000,
        height=700,
        key=f"editor_{document_id}"
    )

    # Handle content updates from the editor
    if result is not None:
        # Check if this is a comment action
        if isinstance(result, dict) and result.get('type') == 'comment_action':
            action = result.get('action')
            print(f"🎯 Received comment action from editor component: {action.get('type') if action else 'None'}")
            
            if action:
                success = _handle_comment_action(db, version_id, action)
                if success:
                    print("✅ Comment action processed successfully, triggering rerun")
                    time.sleep(0.1)
                    st.rerun()
                else:
                    print("❌ Comment action processing failed")
                    st.error("Failed to process comment action. Please try again.")
        
        # Otherwise, it's content update from component
        elif isinstance(result, str):
            print(f"📄 Editor returned content: {len(result)} chars")
            
            # CRITICAL FIX: Always trust component's content - it's the source of truth
            st.session_state.last_component_content = result
            
            # Update editor_content for "Save Version" button to use
            # But DON'T send this back to component
            st.session_state.editor_content = result
            
            # Update original reference if content is growing
            if len(result) > len(st.session_state.get("editor_original_content", "")):
                st.session_state.editor_original_content = result
            
            now = time.time()
            last_save = st.session_state.get("last_autosave", 0)
            
            # Trigger autosave if enough time has passed
            if ((now - last_save) > AUTOSAVE_INTERVAL_SECONDS and 
                not st.session_state.autosave_in_progress):
                
                print(f"💾 Triggering background autosave for {len(result)} chars")
                _handle_progressive_save(db, document_id, result)


# ============================================================================
# AI CHAT HANDLERS
# ============================================================================

def _handle_chat_message(
    db: DatabaseManager,
    version_id: str,
    document_id: str,
    session_id: str,
    message: str,
    document_content: str,
    document_metadata: Dict[str, Any],
    matter_metadata: Dict[str, Any]
) -> Generator[str, None, None]:
    """
    Handle AI chat message and stream response.
    
    Args:
        db: Database manager
        version_id: Current document version ID
        document_id: Document ID
        session_id: Chat session ID
        message: User's message
        document_content: Current document HTML content
        document_metadata: Document metadata
        matter_metadata: Matter metadata
        
    Yields:
        Response chunks from AI
    """
    try:
        from ai_chat_service import AIChatService
        
        # Initialize chat service
        chat_service = AIChatService()
        
        # Save user message to database
        db.create_chat_message(
            version_id=version_id,
            session_id=session_id,
            role='user',
            content=message
        )
        
        # Get conversation history
        chat_history = db.get_chat_history(version_id, session_id, limit=10)
        formatted_history = chat_service.format_conversation_history(chat_history)
        
        # Extract document context
        document_context = chat_service.extract_document_context(
            document=document_metadata,
            matter=matter_metadata,
            current_content=document_content
        )
        
        # Stream AI response
        full_response = ""
        for chunk in chat_service.stream_chat_response(
            user_message=message,
            document_context=document_context,
            conversation_history=formatted_history
        ):
            full_response += chunk
            yield chunk
        
        # Parse any edit suggestions
        edits = chat_service.parse_edit_suggestions(full_response)
        
        # Save AI response to database
        db.create_chat_message(
            version_id=version_id,
            session_id=session_id,
            role='assistant',
            content=full_response,
            metadata={'edits': edits} if edits else None
        )
        
        logger.info(f"Chat message processed with {len(edits)} edit suggestions")
        
    except Exception as e:
        logger.error(f"Error handling chat message: {e}")
        import traceback
        logger.error(traceback.format_exc())
        yield f"\n\n[Error: {str(e)}]"


def _apply_chat_edit(
    db: DatabaseManager,
    document_id: str,
    version_id: str,
    edit: Dict[str, str],
    current_content: str
) -> Tuple[str, bool]:
    """
    Apply an edit suggestion from AI chat.
    
    Args:
        db: Database manager
        document_id: Document ID
        version_id: Current version ID
        edit: Edit specification
        current_content: Current document content
        
    Returns:
        Tuple of (modified_content, success)
    """
    try:
        from ai_chat_service import AIChatService
        
        chat_service = AIChatService()
        modified_content, success = chat_service.apply_edit_to_content(
            content=current_content,
            edit=edit
        )
        
        if success:
            # Update the version content
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(modified_content, 'html.parser')
            content_plain = soup.get_text()
            
            db.update_version_content(
                version_id=version_id,
                content=modified_content,
                content_plain=content_plain
            )
            
            logger.info(f"Successfully applied chat edit to version {version_id}")
            return modified_content, True
        else:
            logger.warning("Failed to apply chat edit")
            return current_content, False
            
    except Exception as e:
        logger.error(f"Error applying chat edit: {e}")
        return current_content, False


# ============================================================================
# MAIN RENDER FUNCTION (continued)
# ============================================================================

    """
    Robust document ID retrieval with multiple fallback strategies.
    Priority order:
    1. Query parameters
    2. Session state (preserved from query params)
    3. Session state (from editor context)
    """
    document_id = None
    
    # Strategy 1: Try query parameters first
    try:
        document_id = st.query_params.get("document_id")
        if document_id:
            logger.info(f"✅ Document ID from query params: {document_id}")
            # Store in session state for future reference
            st.session_state["current_document_id"] = document_id
            st.session_state["preserved_document_id"] = document_id
            return document_id
    except Exception as e:
        logger.warning(f"⚠️ Could not get document_id from query params: {e}")
    
    # Strategy 2: Try preserved session state (from navigation)
    if "preserved_document_id" in st.session_state:
        document_id = st.session_state["preserved_document_id"]
        logger.info(f"✅ Document ID from preserved session state: {document_id}")
        # Restore to query params
        try:
            st.query_params["document_id"] = document_id
        except:
            pass
        return document_id
    
    # Strategy 3: Try current document ID in session state
    if "current_document_id" in st.session_state:
        document_id = st.session_state["current_document_id"]
        logger.info(f"✅ Document ID from current session state: {document_id}")
        # Restore to query params
        try:
            st.query_params["document_id"] = document_id
        except:
            pass
        return document_id
    
    logger.error("❌ No document ID found in any storage location")
    return None

# ============================================================================
# AI CHAT PANEL RENDER FUNCTION
# ============================================================================

@st.fragment
def _render_chat_panel(
    db: DatabaseManager,
    document: Dict[str, Any],
    matter: Dict[str, Any],
    version_id: str,
    document_content: str
):
    """Render the AI chat panel."""
    import uuid
    
    # Initialize chat session
    if "chat_session_id" not in st.session_state:
        # Try to get latest session or create new one
        latest_session = db.get_latest_chat_session(version_id)
        st.session_state.chat_session_id = latest_session or str(uuid.uuid4())
    
    if "chat_messages" not in st.session_state:
        # Load existing messages
        st.session_state.chat_messages = db.get_chat_history(
            version_id,
            st.session_state.chat_session_id,
            limit=50
        )
    
    if "chat_is_streaming" not in st.session_state:
        st.session_state.chat_is_streaming = False
    
    # Chat panel header
    st.markdown("""
    <div style="background: #252930; padding: 12px; border-radius: 8px 8px 0 0; border-bottom: 1px solid #2D3139;">
        <div style="display: flex; align-items: center; gap: 8px;">
            <span style="font-size: 14px; font-weight: 600; color: #FFFFFF;">ClauseBot</span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Messages container
    st.markdown("""
    <style>
    .chat-container {
        background: #1A1D23;
        border-radius: 0 0 8px 8px;
        height: 500px;
        overflow-y: auto;
        padding: 12px;
    }
    .chat-message {
        margin-bottom: 12px;
        padding: 10px 12px;
        border-radius: 8px;
        font-size: 13px;
        line-height: 1.5;
    }
    .chat-message.user {
        background: #4A9EFF;
        color: #FFFFFF;
        margin-left: 20%;
    }
    .chat-message.assistant {
        background: #252930;
        color: #E8EAED;
        border: 1px solid #2D3139;
        margin-right: 20%;
    }
    .chat-role {
        font-size: 11px;
        font-weight: 600;
        margin-bottom: 4px;
        opacity: 0.8;
    }
    .quick-actions-grid {
        display: grid;
        grid-template-columns: repeat(2, 1fr);
        gap: 6px;
        margin: 12px 0;
    }
    /* Target Streamlit buttons inside the chat panel to make them smaller */
    div[data-testid="stVerticalBlock"] > div[data-testid="stHorizontalBlock"] button {
        padding: 0.25rem 0.5rem !important;
        font-size: 0.8rem !important;
        min-height: 0px !important;
        height: auto !important;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Messages display
    messages_placeholder = st.container()
    with messages_placeholder:
        if len(st.session_state.chat_messages) == 0:
            st.markdown("""
            <div style="text-align: center; padding: 40px 20px; color: #9BA1B0;">
                <div style="font-size: 32px; margin-bottom: 12px;">👋</div>
                <p style="font-size: 13px;">Hi! I'm your AI legal assistant.</p>
                <p style="font-size: 12px; opacity: 0.8;">I can help you edit, analyze, and improve this document.</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            for msg in st.session_state.chat_messages:
                if msg['role'] != 'system':
                    role_class = "user" if msg['role'] == "user" else "assistant"
                    role_label = "You" if msg['role'] == "user" else "ClauseBot"
                    
                    st.markdown(f"""
                    <div class="chat-message {role_class}">
                        <div class="chat-role">{role_label}</div>
                        <div>{msg['content']}</div>
                    </div>
                    """, unsafe_allow_html=True)
        
        if st.session_state.chat_is_streaming:
            st.markdown("""
            <div style="padding: 12px; color: #4A9EFF;">
                <em>AI is thinking...</em>
            </div>
            """, unsafe_allow_html=True)
    
    # Quick actions
    st.markdown("<div class='quick-actions-grid'>", unsafe_allow_html=True)
    
    col_q1, col_q2 = st.columns(2)
    with col_q1:
        if st.button("Summarize", use_container_width=True, key="qa_summarize"):
            st.session_state.pending_message = "Provide a brief summary of this document in 2-3 sentences."
            st.rerun()
        if st.button("Citations", use_container_width=True, key="qa_citations"):
            st.session_state.pending_message = "Add relevant Kenyan legal citations to this document."
            st.rerun()
        if st.button("Clarity", use_container_width=True, key="qa_clarity"):
            st.session_state.pending_message = "Identify sections that could be clearer and suggest improvements."
            st.rerun()
    
    with col_q2:
        if st.button("Formalize", use_container_width=True, key="qa_formalize"):
            st.session_state.pending_message = "Make this document more formal and professional."
            st.rerun()
        if st.button("Compliance", use_container_width=True, key="qa_compliance"):
            st.session_state.pending_message = "Check this document for compliance with Kenyan law."
            st.rerun()
        if st.button("Missing Info", use_container_width=True, key="qa_missing"):
            st.session_state.pending_message = "Find any placeholders or missing information."
            st.rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Chat input
    user_input = st.text_area(
        "Message",
        placeholder="Ask a question or request an edit...",
        height=80,
        key="chat_input",
        disabled=st.session_state.chat_is_streaming,
        label_visibility="collapsed"
    )
    
    col_send_clear = st.columns(1)[0]
    with col_send_clear:
        send_button = st.button(
            "Send",
            use_container_width=True,
            type="primary",
            disabled=st.session_state.chat_is_streaming or not user_input.strip()
        )
        
        if st.button("Clear", use_container_width=True):
            st.session_state.chat_messages = []
            st.session_state.chat_session_id = str(uuid.uuid4())
            st.rerun()
    
    # Handle pending quick action message
    if "pending_message" in st.session_state:
        user_input = st.session_state.pending_message
        del st.session_state.pending_message
        send_button = True
    
    # Handle send
    if send_button and user_input.strip():
        st.session_state.chat_is_streaming = True
        
        # Prepare metadata
        doc_metadata = {
            "type": document.get("document_type", "Unknown"),
            "subtype": document.get("document_subtype", ""),
            "title": document.get("title", "Untitled")
        }
        
        matter_metadata = {
            "name": matter.get("name", "Unknown"),
            "client_name": matter.get("client_name", "Unknown"),
            "jurisdiction": matter.get("jurisdiction", "Kenya")
        }
        
        # Stream AI response
        try:
            full_response = ""
            response_placeholder = st.empty()
            
            for chunk in _handle_chat_message(
                db=db,
                version_id=version_id,
                document_id=document['id'],
                session_id=st.session_state.chat_session_id,
                message=user_input.strip(),
                document_content=document_content,
                document_metadata=doc_metadata,
                matter_metadata=matter_metadata
            ):
                full_response += chunk
                response_placeholder.markdown(f"""
                <div class="chat-message assistant">
                    <div class="chat-role">ClauseBot</div>
                    <div>{full_response}</div>
                </div>
                """, unsafe_allow_html=True)
            
            # Reload messages
            st.session_state.chat_messages = db.get_chat_history(
                version_id,
                st.session_state.chat_session_id,
                limit=50
            )
            
        except Exception as e:
            st.error(f"Error: {str(e)}")
        finally:
            st.session_state.chat_is_streaming = False
            st.rerun()


# ============================================================================
# MAIN RENDER FUNCTION
# ============================================================================

def render_document_editor():
    """Render the complete document editor interface with loading animation."""
    
    _init_editor_state()
    
    # CRITICAL: Check if user is authenticated FIRST
    if "user_id" not in st.session_state or not st.session_state.get("user_id"):
        st.error("⚠️ Session expired. Please log in again.")
        st.markdown("""
        <div style="padding: 40px; text-align: center;">
            <div style="font-size: 48px; margin-bottom: 20px;">🔒</div>
            <h3 style="color: #FFFFFF; margin-bottom: 12px;">Authentication Required</h3>
            <p style="color: #9BA1B0; margin-bottom: 24px;">
                Your session has expired. Please log in to continue.
            </p>
            <a href="?view=login" class="sc-btn sc-btn-primary" style="display: inline-block;">
                🔐 Log In
            </a>
        </div>
        """, unsafe_allow_html=True)
        return
    
    db = DatabaseManager()
    db.set_user(st.session_state.user_id)
    
    # CRITICAL FIX: Use robust document ID retrieval with query param fallback
    document_id = get_document_id()

    # Get session param for links to preserve authentication
    session_param = f"&session={st.session_state.get('session_cookie', '')}" if st.session_state.get('session_cookie') else ""
    
    if not document_id:
        st.error("No document ID found")
        st.markdown(f"""
        <div style="padding: 40px; text-align: center;">
            <div style="font-size: 48px; margin-bottom: 20px;">📄</div>
            <h3 style="color: #FFFFFF; margin-bottom: 12px;">Document Not Found</h3>
            <p style="color: #9BA1B0; margin-bottom: 24px;">
                The document you're looking for could not be loaded. This may happen if:<br/>
                • The page was refreshed without proper navigation<br/>
                • The document link is invalid or expired<br/>
                • The document was deleted
            </p>
            <a href="?view=matters{session_param}" target="_self" class="sc-btn sc-btn-primary" style="display: inline-block;">
                ← Back to Matters
            </a>
        </div>
        """, unsafe_allow_html=True)
        return
    
    # Store document ID in session state for persistence within the session
    st.session_state["current_document_id"] = document_id
    st.session_state["preserved_document_id"] = document_id
    
    # ========================================================================
    # FIXED LOADING ANIMATION - Now properly renders before data fetch
    # ========================================================================
    
    # Check if this is the first load (data not yet cached)
    cache_key = f"editor_data_loaded_{document_id}"
    is_first_load = cache_key not in st.session_state
    
    if is_first_load:
        # Create a placeholder for loading animation
        loading_placeholder = st.empty()
        
        # Render the loading animation using the function
        with loading_placeholder.container():
            _render_loading_animation()
        
        # Force a small delay to ensure the animation renders
        time.sleep(0.2)
        
        # Now fetch the data while loading animation is visible
        try:
            # OPTIMIZATION: Use light_mode=True to fetch only metadata initially
            clauses_data = _get_clauses_for_editor(db, light_mode=True)
            document = db.get_document(document_id)
            
            if not document:
                loading_placeholder.empty()
                st.error(f"Document with ID {document_id} not found in database")
                st.markdown(f"""
                <div style="padding: 40px; text-align: center;">
                    <div style="font-size: 48px; margin-bottom: 20px;">🔍</div>
                    <h3 style="color: #FFFFFF; margin-bottom: 12px;">Document Not Found</h3>
                    <p style="color: #9BA1B0; margin-bottom: 24px;">
                        This document may have been deleted or you may not have permission to access it.
                    </p>
                    <a href="?view=matters{session_param}" target="_self" class="sc-btn sc-btn-primary" style="display: inline-block;">
                        ← Back to Matters
                    </a>
                </div>
                """, unsafe_allow_html=True)
                return
            
            matter_id = document['matter_id']
            matter = db.get_matter(matter_id)
            
            if not matter:
                loading_placeholder.empty()
                st.error("Matter not found")
                return
            
            latest_version = db.get_latest_version(document_id)
            
            # Cache the loaded data
            st.session_state[cache_key] = {
                'clauses_data': clauses_data,
                'document': document,
                'matter': matter,
                'latest_version': latest_version
            }
            
            # Clear loading animation
            loading_placeholder.empty()
            
        except Exception as e:
            loading_placeholder.empty()
            st.error(f"Error loading document: {str(e)}")
            logger.error(f"Document loading error: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return
    else:
        # Use cached data (no loading animation on subsequent renders)
        cached_data = st.session_state[cache_key]
        clauses_data = cached_data['clauses_data']
        document = cached_data['document']
        matter = cached_data['matter']
        latest_version = cached_data['latest_version']
        matter_id = document['matter_id']
    
    # CRITICAL: Only load version content on first load
    if st.session_state.generation_complete and latest_version:
        if not st.session_state.get("current_version_id"):
            st.session_state.current_version_id = latest_version["id"]
            st.session_state.editor_content = latest_version["content"]
            st.session_state.editor_original_content = latest_version["content"]
            st.session_state.last_component_content = latest_version["content"]
            # OPTIMIZATION: Initialize comments as empty - will load after editor renders
            if "editor_comments" not in st.session_state:
                st.session_state.editor_comments = []
    
    # ========================================================================
    # GENERATION PHASE
    # ========================================================================
    
    if not st.session_state.generation_complete:
        # CRITICAL: Check if we're actually loading an existing document
        # If document_id exists and we have a version, skip generation entirely
        if document_id and latest_version:
            logger.info(f"✅ Loading existing document {document_id} - skipping generation")
            st.session_state.generation_complete = True
            st.session_state.current_version_id = latest_version["id"]
            st.session_state.editor_content = latest_version["content"]
            st.session_state.editor_original_content = latest_version["content"]
            st.session_state.last_component_content = latest_version["content"]
            # OPTIMIZATION: Initialize comments as empty - will load after editor renders
            if "editor_comments" not in st.session_state:
                st.session_state.editor_comments = []
            # Force a rerun to show the editor
            st.rerun()
            return
        
        # Only show error if truly no data exists
        if "new_matter_payload" not in st.session_state:
            st.error("No document data found. Please create a new matter first.")
            st.markdown(
                f'<a href="?view=matters{session_param}" target="_self" class="sc-btn sc-btn-primary" style="display:inline-block; margin-top:16px;">Back to Matters</a>',
                unsafe_allow_html=True
            )
            return
        
        payload = st.session_state.new_matter_payload
        doc_info = payload.get("document", {})
        
        st.markdown(
            f'''
            <div class="sc-main-header">
                <div class="sc-header-left">
                    <div class="sc-page-title">Generating Document...</div>
                    <div class="sc-page-subtitle">{matter["name"]} • {doc_info.get("subtype") or doc_info.get("type")}</div>
                </div>
            </div>
            ''',
            unsafe_allow_html=True
        )
        
        generator = DocumentGenerator()
        
        with st.spinner("AI is drafting your document..."):
            stream_container = st.empty()
            generated_content = ""
            generation_failed = False
            error_message = ""
            
            # Debounce settings to reduce flashing
            last_update_time = 0
            update_interval = 0.5  # Update UI every 500ms instead of every chunk
            
            try:
                chunk_count = 0
                for chunk in generator.generate_document_stream(payload):
                    generated_content += chunk
                    chunk_count += 1
                    
                    if chunk_count == 1:
                        # First chunk received - generation has started!
                        # Close the modal now so user can see the streaming
                        confirm_generation_started()
                    current_time = time.time()
                    
                    # Only update UI periodically to prevent flashing
                    should_update = (
                        chunk_count == 1 or  # First chunk
                        (current_time - last_update_time) >= update_interval  # Time-based
                    )
                    
                    if not should_update:
                        continue
                    
                    last_update_time = current_time
                    
                    # 1. Clean content: Remove markdown code block markers
                    clean_content = re.sub(r'```[a-z]*\n?', '', generated_content).strip()
                    clean_content = re.sub(r'\n?```$', '', clean_content).strip()
                    
                    # 2. Create complete HTML document for proper rendering
                    html_document = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <style>
                            body {{
                                background-color: #FFFFFF;
                                color: #000000;
                                padding: 40px;
                                margin: 0;
                                font-family: 'Times New Roman', Times, serif;
                                font-size: 12pt;
                                line-height: 1.5;
                            }}
                            h1, h2, h3, h4, h5, h6 {{
                                font-weight: bold;
                                margin-top: 1em;
                                margin-bottom: 0.5em;
                            }}
                            h1 {{ font-size: 18pt; text-align: center; }}
                            h2 {{ font-size: 16pt; }}
                            h3 {{ font-size: 14pt; }}
                            p {{ margin-bottom: 1em; text-align: justify; }}
                            strong, b {{ font-weight: bold; }}
                            em, i {{ font-style: italic; }}
                            u {{ text-decoration: underline; }}
                            ul, ol {{ margin-left: 2em; margin-bottom: 1em; }}
                            li {{ margin-bottom: 0.5em; }}
                            table {{ width: 100%; border-collapse: collapse; margin-bottom: 1em; }}
                            th, td {{ border: 1px solid #000; padding: 8px; text-align: left; }}
                            th {{ background-color: #f0f0f0; font-weight: bold; }}
                        </style>
                    </head>
                    <body>
                        {clean_content}
                    </body>
                    </html>
                    """
                    
                    # 3. Render using components.html with debouncing to prevent flashing
                    with stream_container.container():
                        st.components.v1.html(
                            html_document,
                            height=600,
                            scrolling=True
                        )
                
                # Final update to show complete content
                clean_content = re.sub(r'```[a-z]*\n?', '', generated_content).strip()
                clean_content = re.sub(r'\n?```$', '', clean_content).strip()
                
                html_document = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="UTF-8">
                    <style>
                        body {{
                            background-color: #FFFFFF;
                            color: #000000;
                            padding: 40px;
                            margin: 0;
                            font-family: 'Times New Roman', Times, serif;
                            font-size: 12pt;
                            line-height: 1.5;
                        }}
                        h1, h2, h3, h4, h5, h6 {{
                            font-weight: bold;
                            margin-top: 1em;
                            margin-bottom: 0.5em;
                        }}
                        h1 {{ font-size: 18pt; text-align: center; }}
                        h2 {{ font-size: 16pt; }}
                        h3 {{ font-size: 14pt; }}
                        p {{ margin-bottom: 1em; text-align: justify; }}
                        strong, b {{ font-weight: bold; }}
                        em, i {{ font-style: italic; }}
                        u {{ text-decoration: underline; }}
                        ul, ol {{ margin-left: 2em; margin-bottom: 1em; }}
                        li {{ margin-bottom: 0.5em; }}
                        table {{ width: 100%; border-collapse: collapse; margin-bottom: 1em; }}
                        th, td {{ border: 1px solid #000; padding: 8px; text-align: left; }}
                        th {{ background-color: #f0f0f0; font-weight: bold; }}
                    </style>
                </head>
                <body>
                    {clean_content}
                </body>
                </html>
                """
                
                with stream_container.container():
                    st.components.v1.html(
                        html_document,
                        height=600,
                        scrolling=True
                    )
                
                if not generated_content or len(generated_content.strip()) < 100:
                    generation_failed = True
                    error_message = "Generated content is too short or empty"
                elif "Error" in generated_content[:200] or "error" in generated_content[:200].lower():
                    generation_failed = True
                    error_message = "AI generation returned an error"
                
            except Exception as e:
                generation_failed = True
                error_message = str(e)
                st.error(f"Generation failed: {error_message}")
            
            if not generation_failed and generated_content:
                try:
                    soup = BeautifulSoup(generated_content, 'html.parser')
                    content_plain = soup.get_text()
                    
                    version = db.create_version(
                        document_id=document_id,
                        content=generated_content,
                        content_plain=content_plain,
                        label="Initial Generation",
                        is_major_version=True,
                        change_summary="AI-generated initial draft"
                    )
                    
                    db.update_document_status(document_id, "draft")
                    
                    st.session_state.current_version_id = version["id"]
                    st.session_state.editor_content = generated_content
                    st.session_state.editor_original_content = generated_content
                    st.session_state.last_component_content = generated_content
                    st.session_state.generation_complete = True
                    st.session_state.unsaved_changes = False
                    st.session_state.editor_comments = []  # No comments yet
                    
                    st.success("Document generated successfully!")
                    
                    # PAYWALL: Deduct credit
                    from subscription_manager import SubscriptionManager
                    sub_manager = SubscriptionManager(db)
                    sub_manager.deduct_credit(st.session_state.user_id, document_id)
                    
                    time.sleep(1)
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Failed to save generated document: {str(e)}")
                    db.update_document_status(document_id, "failed")
            else:
                st.error(f"Generation failed: {error_message}")
                st.error("Please check your API configuration and try again.")
                db.update_document_status(document_id, "failed")
                
                if st.button("Back to Matter Details"):
                    st.query_params.update({"view": "matter_details", "matter_id": matter_id})
                    st.rerun()
                
                return
    
    # ========================================================================
    # EDITING PHASE
    # ========================================================================
    
    else:
        # PAYWALL CHECK: Editor Access
        from subscription_manager import SubscriptionManager
        sub_manager = SubscriptionManager(db)
        has_edit_access = sub_manager.has_access(st.session_state.user_id, "document_editor")

        col_h1, col_h2 = st.columns([3, 1])
        with col_h1:
            has_unsaved = st.session_state.editor_content != st.session_state.editor_original_content
            unsaved_indicator = ' <span class="sc-unsaved-indicator" style="top: -4px; position: relative;"></span>' if has_unsaved else ''
            
            st.markdown(
                f'''
                <div class="sc-main-header" style="border:none;padding:16px 0;">
                    <div class="sc-header-left">
                        <div class="sc-page-title">{matter["name"]}{unsaved_indicator}</div>
                        <div class="sc-page-subtitle">{document.get("document_subtype") or document.get("document_type")} • {matter["client_name"]}</div>
                    </div>
                </div>
                ''',
                unsafe_allow_html=True
            )
        
        with col_h2:
            st.markdown("<div style='height:20px;'></div>", unsafe_allow_html=True)
            
            # Header buttons row
            col_v, col_c = st.columns(2)
            with col_v:
                if st.button("Versions", use_container_width=True):
                    st.session_state.show_versions_panel = not st.session_state.show_versions_panel
                    st.rerun()
            with col_c:
                # Initialize chat panel state
                if "show_chat_panel" not in st.session_state:
                    st.session_state.show_chat_panel = True
                
                # CHECK ACCESS: ClauseBot
                has_chatbot_access = sub_manager.has_access(st.session_state.user_id, "ai_chatbot")
                
                if has_chatbot_access:
                    if st.button("ClauseBot", use_container_width=True):
                        st.session_state.show_chat_panel = not st.session_state.show_chat_panel
                        st.rerun()
                else:
                    # Render disabled button for trial users
                    st.button("ClauseBot 🔒", use_container_width=True, disabled=True, help="Upgrade to Individual, Team, or Enterprise to access ClauseBot")
                    # Force panel closed if no access
                    if st.session_state.show_chat_panel:
                        st.session_state.show_chat_panel = False
        
        # Layout based on which panels are open
        panels_open = []
        if st.session_state.show_versions_panel:
            panels_open.append("versions")
        if st.session_state.get("show_chat_panel", False):
            panels_open.append("chat")
        
        # Create appropriate column layout
        if len(panels_open) == 2:  # Both panels
            col_versions, col_editor, col_chat = st.columns([1, 2, 1])
        elif "versions" in panels_open:  # Only versions
            col_versions, col_editor = st.columns([1, 3])
            col_chat = None
        elif "chat" in panels_open:  # Only chat
            col_editor, col_chat = st.columns([3, 1])
            col_versions = None
        else:  # No panels
            col_editor = st.container()
            col_versions = None
            col_chat = None
        
        # Render versions panel if needed
        if col_versions and st.session_state.show_versions_panel:
            with col_versions:
                _render_versions_panel(db, document_id, matter_id)
        
        # Render main editor
        with col_editor:
            if has_edit_access:
                # CRITICAL: Render editor as a fragment so it doesn't reload during autosaves
                _render_editor_fragment(db, document_id, clauses_data, st.session_state.current_version_id)
            else:
                 st.warning("🔒 Editing is restricted to Standard Plan users. You can view and export this document.")
                 
                 # Render Read-Only View
                 content_html = st.session_state.editor_content
                 
                 # sanitize slightly for display
                 st.markdown(f"""
                 <div style="border: 1px solid #4B5563; border-radius: 8px; padding: 40px; background: white; color: black; height: 600px; overflow-y: auto; box-shadow: inset 0 2px 4px rgba(0,0,0,0.1);">
                    {content_html}
                 </div>
                 """, unsafe_allow_html=True)
                 


        if st.session_state.show_versions_panel:
            with col_versions:
                _render_versions_panel(db, document_id, st.session_state.current_version_id)
                st.markdown("</div>", unsafe_allow_html=True)
        
        # Render chat panel if needed
        if col_chat and st.session_state.get("show_chat_panel", False):
            with col_chat:
                _render_chat_panel(
                    db=db,
                    document=document,
                    matter=matter,
                    version_id=st.session_state.current_version_id,
                    document_content=st.session_state.editor_content
                )
        
        st.markdown("<div style='height:24px;'></div>", unsafe_allow_html=True)
        
        # --- Action Buttons ---
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button("Save Version", use_container_width=True, type="primary", disabled=not has_edit_access):
                current_content = st.session_state.get("editor_content", "")
                word_count = _word_count(current_content)

                if word_count < 5:
                    st.error(f"Cannot save: Document must contain at least 5 words (found {word_count})")
                    return
                
                with st.spinner("Saving..."):
                    content_plain = _extract_plain_text(current_content)
                    
                    versions = db.get_versions(document_id, include_content=False)
                    next_version_num = len(versions) + 1
                    
                    try:
                        version = db.create_version(
                            document_id=document_id,
                            content=current_content,
                            content_plain=content_plain,
                            label=f"Version {next_version_num}",
                            is_major_version=False,
                            change_summary=None
                        )
                        
                        st.session_state.current_version_id = version["id"]
                        st.session_state.editor_original_content = current_content
                        st.session_state.last_autosave = time.time()
                        
                        # Reload comments for new version
                        st.session_state.editor_comments = _get_comments_for_version(db, version["id"])
                        
                        db.update_document_status(document_id, "draft")
                        
                        st.success("Version saved!")
                        time.sleep(1)
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"Failed to save version: {str(e)}")
                        import traceback
                        st.error(traceback.format_exc())
        
        with col2:
            # Generate DOCX from current editor content
            try:
                import tempfile
                import os
                
                export_content = st.session_state.get("editor_content", "")
                
                if not export_content or len(export_content.strip()) < 10:
                    st.button("Export DOCX", use_container_width=True, disabled=True, 
                             help="No content to export")
                else:
                    version_label = db.get_version(st.session_state.current_version_id).get('label', 'draft')
                    safe_matter_name = matter['name'].replace(' ', '_')
                    safe_version_label = version_label.replace(' ', '_')
                    docx_filename = f"{safe_matter_name}_{safe_version_label}.docx"
                    
                    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                    tmp_file_path = tmp_file.name
                    tmp_file.close()
                    
                    try:
                        _html_to_docx(
                            export_content,
                            tmp_file_path,
                            matter['name']
                        )
                        
                        with open(tmp_file_path, "rb") as f:
                            docx_data = f.read()
                        
                    finally:
                        try:
                            os.unlink(tmp_file_path)
                        except Exception as cleanup_error:
                            print(f"Warning: Could not delete temp file: {cleanup_error}")
                    
                    st.download_button(
                        label="Export DOCX",
                        data=docx_data,
                        file_name=docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="download_docx_btn",
                        help=f"Download current version as {docx_filename}"
                    )
            except Exception as e:
                st.error(f"Export failed: {str(e)}")
                import traceback
                st.error(traceback.format_exc())
        
        with col3:
            if st.button("Save Major Version", use_container_width=True):
                st.session_state.show_major_version_modal = True
        
        with col4:
            if st.button("Back to Matter", use_container_width=True):
                has_unsaved = st.session_state.editor_content != st.session_state.editor_original_content
                
                if has_unsaved:
                    if st.session_state.get("confirm_exit"):
                        st.session_state.confirm_exit = False
                        # Clear cache when navigating away
                        if cache_key in st.session_state:
                            del st.session_state[cache_key]
                        # CRITICAL FIX: Preserve session state when navigating
                        from auth import update_query_params
                        update_query_params({
                            "view": "matter_details", 
                            "matter_id": matter_id
                        })
                        st.rerun()
                    else:
                        st.warning("You have unsaved changes. Click again to confirm exit.")
                        st.session_state.confirm_exit = True
                else:
                    # Clear cache when navigating away
                    if cache_key in st.session_state:
                        del st.session_state[cache_key]
                    # CRITICAL FIX: Preserve session state when navigating
                    from auth import update_query_params
                    update_query_params({
                        "view": "matter_details", 
                        "matter_id": matter_id
                    })
                    st.rerun()
        
        if st.session_state.get("show_major_version_modal", False):
            _render_major_version_modal(db, document_id)


# ============================================================================
# MAJOR VERSION MODAL
# ============================================================================

@st.dialog("Save Major Version")
def _render_major_version_modal(db: DatabaseManager, document_id: str):
    """Render modal for saving a major version."""
    
    st.markdown("""
    <div style="color: #9BA1B0; font-size: 14px; margin-bottom: 20px;">
        Major versions are significant milestones in your document's lifecycle (e.g., "Final Draft", "Client Review", "Executed").
    </div>
    """, unsafe_allow_html=True)
    
    version_label = st.text_input(
        "Version Label*",
        placeholder="e.g., Final Draft, Client Review v1, Executed",
        help="Give this version a meaningful name"
    )
    
    change_summary = st.text_area(
        "Summary of Changes",
        placeholder="Describe what changed in this version...",
        help="Optional: Summarize the key changes made",
        height=100
    )
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Cancel", use_container_width=True):
            st.session_state.show_major_version_modal = False
            st.rerun()
    
    with col2:
        if st.button("Save Major Version", use_container_width=True, type="primary"):
            if not version_label or not version_label.strip():
                st.error("Please provide a version label")
                return
            
            current_content = st.session_state.get("editor_content", "")
            word_count = _word_count(current_content)

            if word_count < 5:
                st.error(f"Cannot save: Document must contain at least 5 words (found {word_count})")
                return
            
            with st.spinner("Saving major version..."):
                content_plain = _extract_plain_text(current_content)
                
                try:
                    version = db.create_version(
                        document_id=document_id,
                        content=current_content,
                        content_plain=content_plain,
                        label=version_label.strip(),
                        is_major_version=True,
                        change_summary=change_summary.strip() if change_summary else None
                    )

                    st.session_state.current_version_id = version["id"]
                    st.session_state.editor_original_content = current_content
                    st.session_state.unsaved_changes = False
                    st.session_state.show_major_version_modal = False
                    st.session_state.last_autosave = time.time()
                    
                    # Reload comments for new version
                    st.session_state.editor_comments = _get_comments_for_version(db, version["id"])
                    
                    st.success(f"Major version '{version_label}' saved!")
                    time.sleep(1)
                    st.rerun()
                    
                except Exception as e:
                    st.error(f"Failed to save major version: {str(e)}")
                    import traceback
                    st.error(traceback.format_exc())